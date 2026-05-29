"""Generate DOCX report from JSON.

Usage: python gen_report_docx.py <report.json> <output.docx>
Reads report JSON from file, writes DOCX to output path.
"""

import json
import sys
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

PRIMARY = RGBColor(0x25, 0x63, 0xEB)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def load_report(path):
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def strip_tags(html):
    """Convert basic HTML to plain text with some formatting preserved."""
    text = html
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'</p>', '\n', text)
    text = re.sub(r'</li>', '\n', text)
    text = re.sub(r'<[^>]*>', '', text)
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&nbsp;', ' ')
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)
    return text.strip()


def create_docx(report, output_path):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    meta = report.get('meta', {})
    sections = report.get('sections', [])
    refs = report.get('references', [])

    # ── Cover Page ──
    # Add some spacing before title
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(report.get('title', '研究报告'))
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = meta.get('generatedAt', '')[:10] if meta.get('generatedAt') else ''
    run = sub_para.add_run(f"{report.get('type', '深度报告')}  |  {date_str}")
    run.font.size = Pt(14)
    run.font.color.rgb = GRAY

    # Stats
    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = []
    if meta.get('sourceCount'):
        parts.append(f"素材来源: {meta['sourceCount']} 条")
    if meta.get('citedCount'):
        parts.append(f"引用: {meta['citedCount']} 条")
    if meta.get('verifiedClaims'):
        parts.append(f"已验证主张: {meta['verifiedClaims']} 条")
    run = stats_para.add_run("  |  ".join(parts))
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    doc.add_page_break()

    # ── Table of Contents ──
    toc_heading = doc.add_heading('目录', level=1)
    for i, s in enumerate(sections):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"{i + 1}. {s['title']}")
        run.font.size = Pt(13)
        run.font.color.rgb = DARK

    doc.add_page_break()

    # ── Content Sections ──
    for section_data in sections:
        doc.add_heading(section_data['title'], level=1)

        text = strip_tags(section_data.get('content', ''))

        # Split into paragraphs
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        for para_text in paragraphs:
            if para_text:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = Pt(22)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(para_text)
                run.font.size = Pt(11)
                run.font.color.rgb = DARK

        # Citation note
        cit_count = len(section_data.get('citations', []))
        if cit_count > 0:
            p = doc.add_paragraph()
            run = p.add_run(f"[本节引用 {cit_count} 条来源]")
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
            run.font.italic = True

    doc.add_page_break()

    # ── References ──
    doc.add_heading('参考文献', level=1)

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        cred = f"[{ref.get('credibility', 'B')}级]"
        freshness = ref.get('freshness', '')
        freshness_label = '有效' if freshness == 'green' else '偏旧' if freshness == 'yellow' else '过时' if freshness == 'red' else ''

        run = p.add_run(f"[{ref['index']}] ")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

        run2 = p.add_run(ref['title'])
        run2.font.size = Pt(9)
        run2.font.color.rgb = DARK

        meta_parts = [cred]
        if freshness_label:
            meta_parts.append(freshness_label)
        if ref.get('publishedAt'):
            meta_parts.append(ref['publishedAt'][:10])

        run3 = p.add_run(f"  {' · '.join(meta_parts)}")
        run3.font.size = Pt(8)
        run3.font.color.rgb = GRAY

        if ref.get('url'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(8)
            run4 = p2.add_run(ref['url'])
            run4.font.size = Pt(8)
            run4.font.color.rgb = PRIMARY

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python gen_report_docx.py <report.json> <output.docx>")
        sys.exit(1)

    report_path = sys.argv[1]
    output_path = sys.argv[2]

    report = load_report(report_path)
    result = create_docx(report, output_path)
    print(f"DOCX saved: {result}")
